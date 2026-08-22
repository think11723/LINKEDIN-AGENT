"""File extraction — Phase 10.

Safe text extraction for PDF and DOCX uploads. The output is plain
text suitable for the deterministic resume parser.

Hard rules:

* Files larger than the limit are rejected before extraction
  (memory safety).
* Only PDF and DOCX MIME types / extensions are accepted.
* Filenames are sanitized; user-controlled paths are never used.
* The PDF library (``pypdf``) and the DOCX library
  (``python-docx``) are run in-process; both are well-tested
  open-source projects. They have no C extensions and no
  network access.
"""

from __future__ import annotations

import io
import re
from pathlib import PurePosixPath
from typing import Tuple


# Hard cap on uploaded file size. 10 MB is enough for any real
# resume and prevents memory abuse.
MAX_FILE_BYTES = 10 * 1024 * 1024

# Allowed content types and extensions.
ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    # Some browsers / OSes report DOCX as one of these:
    "application/zip",
    "application/octet-stream",
}
ALLOWED_EXTENSIONS = {"pdf", "docx"}


class FileExtractionError(Exception):
    """Raised on any user-facing extraction failure.

    The HTTP layer should surface a user-safe message that does
    NOT echo the filename or internal state.
    """

    def __init__(self, message: str, *, code: str = "extraction_failed") -> None:
        super().__init__(message)
        self.message = message
        self.code = code


def _sanitize_filename(filename: str) -> str:
    """Return a safe filename. User input is untrusted; we strip
    path components and any control characters."""
    if not filename:
        return "upload"
    # PurePosixPath strips any path components and returns only
    # the final segment.
    base = PurePosixPath(filename).name
    # Drop Windows drive letters, etc.
    base = base.replace(":", "").replace("\\", "_")
    base = re.sub(r"[\x00-\x1f\x7f]", "", base)
    return base[:120] or "upload"


def validate_upload(content_type: str | None, filename: str | None) -> Tuple[str, str]:
    """Validate MIME type + filename. Return ``(content_type, safe_filename)``
    or raise :class:`FileExtractionError` with a user-safe message.
    """
    safe_name = _sanitize_filename(filename or "upload")
    ext = safe_name.rsplit(".", 1)[-1].lower() if "." in safe_name else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise FileExtractionError(
            "Unsupported file type. Please upload a PDF or DOCX.",
            code="unsupported_format",
        )
    ct = (content_type or "").lower().split(";")[0].strip()
    if ct and ct not in ALLOWED_CONTENT_TYPES:
        raise FileExtractionError(
            "Unsupported file type. Please upload a PDF or DOCX.",
            code="unsupported_format",
        )
    if not ct:
        # Some clients send empty content-type. Use the extension
        # to derive a default.
        ct = (
            "application/pdf"
            if ext == "pdf"
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    return ct, safe_name


def _extract_pdf(data: bytes) -> str:
    """Extract text from a PDF using pypdf.

    pypdf is intentionally pure-Python. We iterate over pages
    and concatenate the extracted text.
    """
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as e:  # pragma: no cover
        raise FileExtractionError(
            "PDF parsing is not available on this server.",
            code="parser_unavailable",
        ) from e
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:  # pypdf raises a wide variety of types
        raise FileExtractionError(
            "This PDF could not be read. It may be password-protected or corrupt.",
            code="bad_pdf",
        ) from e
    parts: list = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    """Extract text from a DOCX using python-docx."""
    try:
        import docx  # type: ignore
    except ImportError as e:  # pragma: no cover
        raise FileExtractionError(
            "DOCX parsing is not available on this server.",
            code="parser_unavailable",
        ) from e
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as e:
        raise FileExtractionError(
            "This DOCX could not be read. It may be password-protected or corrupt.",
            code="bad_docx",
        ) from e
    parts: list = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)
    # Also iterate tables to catch tabular resumes.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if text:
                    parts.append(text)
    return "\n".join(parts)


def extract_text(content_type: str, filename: str, data: bytes) -> Tuple[str, str]:
    """Validate the upload and extract plain text.

    Returns ``(text, safe_filename)``. Raises
    :class:`FileExtractionError` on any failure with a user-safe
    message.
    """
    if not data:
        raise FileExtractionError("Uploaded file is empty.", code="empty_file")
    if len(data) > MAX_FILE_BYTES:
        raise FileExtractionError(
            f"File is too large. Maximum size is {MAX_FILE_BYTES // (1024 * 1024)} MB.",
            code="file_too_large",
        )
    ct, safe_name = validate_upload(content_type, filename)
    ext = safe_name.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        text = _extract_pdf(data)
    elif ext == "docx":
        text = _extract_docx(data)
    else:
        # validate_upload should have caught this; defensive.
        raise FileExtractionError(
            "Unsupported file type. Please upload a PDF or DOCX.",
            code="unsupported_format",
        )
    # Cap the extracted text at a sensible upper bound so a
    # pathological PDF cannot produce gigabytes of text.
    if len(text) > 500_000:
        text = text[:500_000]
    return text, safe_name
